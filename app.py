import streamlit as st
import zipfile
import xml.etree.ElementTree as ET
import os
import fitz  # PyMuPDF
from docx import Document
import tempfile
from PIL import Image
import io
import re
from datetime import datetime


def analyze_pdf(file_bytes):
    metadata = {}
    with fitz.open(stream=file_bytes, filetype="pdf") as pdf:
        metadata['basic'] = pdf.metadata

        # Extract full document content
        full_text = ""
        metadata['content'] = {
            'annotations': {},
            'links': {},
            'forms': {}
        }

        # Collect all mailto links
        mailto_links = []

        for page_num in range(len(pdf)):
            page = pdf[page_num]

            # Get text for analysis
            full_text += page.get_text()

            # Get annotations
            annotations = page.annots()
            if annotations:
                metadata['content']['annotations'][f'page_{page_num+1}'] = [
                    {
                        'type': annot.type[1],
                        'content': annot.info.get('content', ''),
                        'modified': annot.info.get('modified', ''),
                        'subject': annot.info.get('subject', '')
                    } for annot in annotations
                ]

            # Get links and collect mailto links
            links = page.get_links()
            if links:
                metadata['content']['links'][f'page_{page_num+1}'] = []
                for link in links:
                    link_data = {
                        'type': link.get('type', ''),
                        'uri': link.get('uri', ''),
                        'destination': link.get('destination', '')
                    }
                    metadata['content']['links'][f'page_{page_num+1}'].append(
                        link_data)

                    # If it's a mailto link, add to sensitive data
                    if link_data['uri'] and link_data['uri'].startswith('mailto:'):
                        mailto_links.append(link_data['uri'])

            # Get form fields
            widgets = page.widgets()
            if widgets:
                metadata['content']['forms'][f'page_{page_num+1}'] = [
                    {
                        'field_type': widget.field_type,
                        'field_name': widget.field_name,
                        'value': widget.field_value,
                    } for widget in widgets
                ]

        # Extract images and their metadata
        metadata['images'] = {}
        for page_num in range(len(pdf)):
            page = pdf[page_num]
            images = page.get_images()

            for img_index, img in enumerate(images):
                xref = img[0]
                base_image = pdf.extract_image(xref)
                if base_image:
                    image_bytes = base_image["image"]
                    metadata['images'][f'page_{page_num}_img_{img_index}'] = get_image_metadata(
                        image_bytes)

        # Analyze all text for sensitive data
        metadata['sensitive_data'] = find_sensitive_patterns(full_text)

        # Add mailto links to sensitive data if any were found
        if mailto_links:
            if 'emails_and_mailtos' not in metadata['sensitive_data']:
                metadata['sensitive_data']['emails_and_mailtos'] = []
            metadata['sensitive_data']['emails_and_mailtos'].extend(
                mailto_links)
            # Remove duplicates while preserving order
            metadata['sensitive_data']['emails_and_mailtos'] = list(dict.fromkeys(
                metadata['sensitive_data']['emails_and_mailtos']
            ))

    return metadata


def analyze_docx(file_bytes):
    metadata = {}

    with tempfile.NamedTemporaryFile(delete=False) as tmp:
        tmp.write(file_bytes)
        tmp_path = tmp.name

    try:
        # Get core document properties
        doc = Document(tmp_path)

        # Get ZIP info for detailed metadata
        with zipfile.ZipFile(tmp_path, 'r') as zip_ref:
            metadata['file_metadata'] = {
                'MIME_Type': 'application/vnd.openxmlformats-officedocument.wordprocessingml.document',
                'Application': doc.core_properties.identifier or '',
                'Doc_Security': 'None',  # Default unless otherwise specified
                'Create_Date': str(doc.core_properties.created) if doc.core_properties.created else '',
                'Creator': doc.core_properties.author or '',
                'Last_Modified_By': doc.core_properties.last_modified_by or '',
                'Modify_Date': str(doc.core_properties.modified) if doc.core_properties.modified else '',
                'Created': str(doc.core_properties.created) if doc.core_properties.created else '',
                'LastSaved': str(doc.core_properties.modified) if doc.core_properties.modified else '',
            }

            # Try to get additional metadata from app.xml
            if 'docProps/app.xml' in zip_ref.namelist():
                with zip_ref.open('docProps/app.xml') as app_xml:
                    tree = ET.parse(app_xml)
                    root = tree.getroot()

                    # Extract all properties from app.xml
                    for elem in root:
                        tag = elem.tag.split('}')[-1]
                        if elem.text and tag not in ['Template', 'Pages', 'Words', 'Characters', 'CharactersWithSpaces']:
                            metadata['file_metadata'][tag] = elem.text

        # Analyze ZIP contents
        with zipfile.ZipFile(tmp_path, 'r') as zip_ref:
            metadata['zip_contents'] = {}

            # List all files in the ZIP
            all_files = zip_ref.namelist()
            metadata['zip_contents']['files'] = all_files

            # Get document statistics
            metadata['statistics'] = {
                'total_files': len(all_files),
                'media_files': len([f for f in all_files if 'media/' in f]),
                'embedded_files': len([f for f in all_files if 'embeddings/' in f]),
                'custom_xml': len([f for f in all_files if 'customXml/' in f])
            }

            # Improved XML text extraction
            all_text = []
            for xml_file in [f for f in all_files if f.endswith('.xml')]:
                try:
                    with zip_ref.open(xml_file) as f:
                        content = f.read().decode('utf-8')
                        # Remove XML tags but keep content
                        text_content = re.sub(r'<[^>]+>', ' ', content)
                        # Clean up whitespace
                        text_content = re.sub(
                            r'\s+', ' ', text_content).strip()
                        all_text.append(text_content)
                except Exception as e:
                    continue

            # Get visible document content
            visible_text = '\n'.join([
                para.text for para in doc.paragraphs if para.text.strip()
            ])
            all_text.append(visible_text)

            # Tables content
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            all_text.append(cell.text)

            # Analyze all text content together
            full_text = ' '.join(all_text)
            # Clean up the text to help with pattern matching
            full_text = re.sub(r'\s+', ' ', full_text)
            metadata['sensitive_data'] = find_sensitive_patterns(full_text)

            # Extract and analyze embedded images
            metadata['embedded_images'] = {}
            for file in all_files:
                if any(file.startswith(prefix) for prefix in ['word/media/', 'word/embeddings/']):
                    try:
                        with zip_ref.open(file) as img_file:
                            img_bytes = img_file.read()
                            metadata['embedded_images'][file] = get_image_metadata(
                                img_bytes)
                    except Exception:
                        continue

    finally:
        os.unlink(tmp_path)

    return metadata


def get_image_metadata(image_bytes):
    try:
        img = Image.open(io.BytesIO(image_bytes))
        return {
            'format': img.format,
            'size': img.size,
            'mode': img.mode,
            'exif': img.getexif().get_ifd(0x8825) if hasattr(img, '_getexif') else None,
        }
    except Exception as e:
        return str(e)


def find_sensitive_patterns(text):
    patterns = {
        # Even more robust email pattern
        'emails_and_mailtos': r'(?:mailto:)?(?:[a-zA-Z0-9!#$%&\'*+/=?^_`{|}~-]+(?:\.[a-zA-Z0-9!#$%&\'*+/=?^_`{|}~-]+)*|"(?:[\x01-\x08\x0b\x0c\x0e-\x1f\x21\x23-\x5b\x5d-\x7f]|\\[\x01-\x09\x0b\x0c\x0e-\x7f])*")@(?:(?:[a-zA-Z0-9](?:[a-zA-Z0-9-]*[a-zA-Z0-9])?\.)+[a-zA-Z0-9](?:[a-zA-Z0-9-]*[a-zA-Z0-9])?|\[(?:(?:25[0-5]|2[0-4][0-9]|[01]?[0-9][0-9]?)\.){3}(?:25[0-5]|2[0-4][0-9]|[01]?[0-9][0-9]?|[a-zA-Z0-9-]*[a-zA-Z0-9]:(?:[\x01-\x08\x0b\x0c\x0e-\x1f\x21-\x5a\x53-\x7f]|\\[\x01-\x09\x0b\x0c\x0e-\x7f])+)\])',
        'ip_addresses': r'\b(?:\d{1,3}\.){3}\d{1,3}\b',
        'network_paths': r'\\\\[a-zA-Z0-9\-_]+\\[a-zA-Z0-9\-_$]+(?:\\[a-zA-Z0-9\-_$]+)*',
        'local_paths': r'[A-Za-z]:\\[^<>:"|?*\n\r]+',
        'usernames': r'(?i)user(?:name)?["\'\s:]+([a-zA-Z0-9\-_@\.]+)',
        'api_keys': r'(?i)(api[_-]?key|access[_-]?token)["\'\s:]+([a-zA-Z0-9\-_]{16,})',
        'urls': r'https?://(?:[-\w.]|(?:%[\da-fA-F]{2}))+[^\s]*',
        'social_security': r'\b\d{3}-\d{2}-\d{4}\b',
        'credit_cards': r'\b(?:\d[ -]*?){13,16}\b',
        'phone_numbers': r'\b(?:\+\d{1,2}\s?)?\(?\d{3}\)?[\s.-]?\d{3}[\s.-]?\d{4}\b'
    }

    results = {}
    for pattern_name, pattern in patterns.items():
        matches = re.findall(pattern, text, re.IGNORECASE)
        if matches:
            # Handle both single matches and tuple matches from groups
            cleaned_matches = []
            for match in matches:
                if isinstance(match, tuple):
                    cleaned_matches.extend(m for m in match if m)
                else:
                    cleaned_matches.append(match)
            results[pattern_name] = list(set(cleaned_matches))
    return results


def analyze_security_concerns(metadata, file_type):
    concerns = []

    # Check emails
    if 'sensitive_data' in metadata and 'emails_and_mailtos' in metadata['sensitive_data']:
        emails = metadata['sensitive_data']['emails_and_mailtos']
        if len(emails) > 1:
            concerns.append({
                'level': 'WARNING',
                'message': f'Multiple email addresses found ({len(emails)})',
                'details': emails
            })

        # Check for numbers in emails
        emails_with_numbers = [email for email in emails if any(
            char.isdigit() for char in email)]
        if emails_with_numbers:
            concerns.append({
                'level': 'INFO',
                'message': 'Emails containing numbers detected',
                'details': emails_with_numbers
            })

    # Check software fingerprints in both DOCX and PDF
    if file_type == 'application/vnd.openxmlformats-officedocument.wordprocessingml.document':
        if 'file_metadata' in metadata:
            app_info = str(metadata['file_metadata'].get(
                'Application', '')).lower()
            if 'wps' in app_info:
                concerns.append({
                    'level': 'WARNING',
                    'message': 'Document created with WPS Office',
                    'details': metadata['file_metadata'].get('Application', '')
                })
            if 'kendo' in app_info:
                concerns.append({
                    'level': 'WARNING',
                    'message': 'Document contains Kendo UI traces',
                    'details': metadata['file_metadata'].get('Application', '')
                })
    elif file_type == 'application/pdf':
        if 'basic' in metadata:
            # Check Creator and Producer fields
            creator = str(metadata['basic'].get('creator', '')).lower()
            producer = str(metadata['basic'].get('producer', '')).lower()

            if 'wps' in creator or 'wps' in producer:
                concerns.append({
                    'level': 'WARNING',
                    'message': 'Document created with WPS Office',
                    'details': f"Creator: {metadata['basic'].get('creator', '')}, Producer: {metadata['basic'].get('producer', '')}"
                })
            if 'kendo' in creator or 'kendo' in producer:
                concerns.append({
                    'level': 'WARNING',
                    'message': 'Document contains Kendo UI traces',
                    'details': f"Creator: {metadata['basic'].get('creator', '')}, Producer: {metadata['basic'].get('producer', '')}"
                })

    return concerns


st.title("Document Metadata Analyzer")

uploaded_file = st.file_uploader(
    "Upload a PDF or DOCX file", type=['pdf', 'docx'])

if uploaded_file:
    file_bytes = uploaded_file.read()
    file_type = uploaded_file.type

    st.write("### File Information")
    st.write(f"Filename: {uploaded_file.name}")
    st.write(f"File type: {file_type}")

    try:
        if file_type == 'application/pdf':
            metadata = analyze_pdf(file_bytes)

            concerns = analyze_security_concerns(metadata, file_type)
            if concerns:
                st.write("### ⚠️ Security Concerns")
                for concern in concerns:
                    if concern['level'] == 'WARNING':
                        st.warning(
                            f"**{concern['message']}**\n\n{concern['details']}")
                    else:
                        st.info(
                            f"**{concern['message']}**\n\n{concern['details']}")

            st.write("### PDF Basic Metadata")
            # Highlight WPS or Kendo UI in metadata
            for key, value in metadata.get('basic', {}).items():
                if value and ('wps' in str(value).lower() or 'kendo' in str(value).lower()):
                    st.warning(f"**{key}:** {value}")
                else:
                    st.write(f"**{key}:** {value}")

            if metadata.get('content'):
                if metadata['content']['annotations']:
                    st.write("### Annotations Found")
                    st.json(metadata['content']['annotations'])

                if metadata['content']['links']:
                    st.write("### Links Found")
                    st.json(metadata['content']['links'])

                if metadata['content']['forms']:
                    st.write("### Form Fields Found")
                    st.json(metadata['content']['forms'])

            if metadata.get('images'):
                st.write("### Embedded Images")
                st.json(metadata['images'])

            if metadata.get('sensitive_data'):
                st.write("### Sensitive Data Found")
                st.json(metadata['sensitive_data'])

        elif file_type == 'application/vnd.openxmlformats-officedocument.wordprocessingml.document':
            metadata = analyze_docx(file_bytes)

            # Analyze security concerns
            concerns = analyze_security_concerns(metadata, file_type)
            if concerns:
                st.write("### ⚠️ Security Concerns")
                for concern in concerns:
                    if concern['level'] == 'WARNING':
                        st.warning(
                            f"**{concern['message']}**\n\n{concern['details']}")
                    else:
                        st.info(
                            f"**{concern['message']}**\n\n{concern['details']}")

            st.write("### DOCX File Metadata")
            if metadata.get('file_metadata'):
                # Highlight WPS or Kendo UI in metadata
                for key, value in metadata['file_metadata'].items():
                    if key == 'Application' and value and ('wps' in value.lower() or 'kendo' in value.lower()):
                        st.warning(f"**{key.replace('_', ' ')}:** {value}")
                    else:
                        st.write(f"**{key.replace('_', ' ')}:** {value}")

            if metadata.get('statistics'):
                st.write("### Document Statistics")
                st.json(metadata['statistics'])

            if metadata.get('zip_contents'):
                st.write("### Document Structure")
                st.write("#### Files in Archive")
                st.json(metadata['zip_contents']['files'])

            if metadata.get('sensitive_data'):
                st.write("### Sensitive Data Found")
                st.json(metadata['sensitive_data'])

            if metadata.get('embedded_images'):
                st.write("### Embedded Images Metadata")
                st.json(metadata['embedded_images'])

    except Exception as e:
        st.error(f"Error analyzing file: {str(e)}")
